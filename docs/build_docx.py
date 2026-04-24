"""Generate docs/rapport.docx from the project content + screenshots.

Run via a disposable Python container (no global install):

    docker run --rm -v "$PWD:/app" -w /app python:3.11-slim \
        bash -c "pip install --quiet python-docx && python docs/build_docx.py"
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
SHOTS = ROOT / "docs" / "screenshots"
import os
OUT = ROOT / "docs" / os.environ.get("RAPPORT_FILENAME", "rapport.docx")


INDIGO = RGBColor(0x43, 0x38, 0xCA)
DARK = RGBColor(0x1F, 0x29, 0x37)
GRAY = RGBColor(0x4B, 0x55, 0x63)
MUTED = RGBColor(0x6B, 0x72, 0x80)
CODE_BG = "F3F4F6"


def set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = INDIGO if level <= 1 else DARK
        run.font.name = "Calibri"


def add_paragraph(doc: Document, text: str, *, bold: bool = False,
                   color: RGBColor = DARK, size: int = 11,
                   align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def add_bullet(doc: Document, text: str, runs_bold: tuple[str, ...] = ()) -> None:
    p = doc.add_paragraph(style="List Bullet")
    remaining = text
    for fragment in runs_bold:
        idx = remaining.find(fragment)
        if idx > 0:
            run = p.add_run(remaining[:idx])
            run.font.size = Pt(11)
        if idx >= 0:
            run = p.add_run(fragment)
            run.bold = True
            run.font.size = Pt(11)
            remaining = remaining[idx + len(fragment):]
    run = p.add_run(remaining)
    run.font.size = Pt(11)


def add_code_block(doc: Document, code: str, *, language: str = "") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_bg(cell, CODE_BG)
    cell.width = Cm(16)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    para = cell.paragraphs[0]
    para.paragraph_format.left_indent = Cm(0.2)
    para.paragraph_format.right_indent = Cm(0.2)

    for i, line in enumerate(code.rstrip("\n").split("\n")):
        if i > 0:
            para = cell.add_paragraph()
            para.paragraph_format.left_indent = Cm(0.2)
            para.paragraph_format.right_indent = Cm(0.2)
        run = para.add_run(line or " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = DARK

    doc.add_paragraph()


def add_screenshot(doc: Document, filename: str, caption: str) -> None:
    path = SHOTS / filename
    if not path.exists():
        add_paragraph(doc, f"[Screenshot manquant : {filename}]", color=MUTED)
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(15))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = MUTED


def build() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # =================================================================
    # Cover
    # =================================================================
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_run = cover.add_run("\n\n\nMini-Projet DevOps")
    cover_run.bold = True
    cover_run.font.size = Pt(26)
    cover_run.font.color.rgb = INDIGO

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Automatisation du déploiement d'une application web")
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = DARK
    sub_run.italic = True

    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_run = sep.add_run("\n" + "─" * 40 + "\n")
    sep_run.font.color.rgb = MUTED

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for label, value in [
        ("Application", "FormaPro — plateforme de formations"),
        ("Stack", "Laravel 10 · PHP 8.2 · MySQL 8 · Docker · GitHub Actions"),
        ("Auteur", "Reda Ben Ammi"),
        ("Date", "Avril 2026"),
    ]:
        r1 = info.add_run(f"{label} : ")
        r1.bold = True
        r1.font.size = Pt(12)
        r2 = info.add_run(f"{value}\n")
        r2.font.size = Pt(12)

    doc.add_page_break()

    # =================================================================
    # 1. Introduction
    # =================================================================
    add_heading(doc, "1. Introduction", level=1)
    add_paragraph(
        doc,
        "Ce projet met en œuvre une chaîne d'automatisation DevOps complète pour "
        "une application web Laravel : développement, conteneurisation Docker, "
        "orchestration via Docker Compose et pipeline CI/CD GitHub Actions. "
        "L'application, baptisée FormaPro, permet de consulter une liste de "
        "formations techniques et d'en ajouter de nouvelles via un formulaire.",
    )
    add_paragraph(
        doc,
        "Contrainte technique : aucune dépendance n'est installée globalement sur "
        "la machine hôte. Tous les outils (Composer, Python, Chromium pour les "
        "captures) sont exécutés dans des conteneurs Docker jetables.",
    )

    # =================================================================
    # 2. Architecture
    # =================================================================
    add_heading(doc, "2. Description de l'architecture", level=1)

    add_heading(doc, "2.1 Vue d'ensemble", level=2)
    add_paragraph(
        doc,
        "L'application expose cinq routes :",
    )
    add_bullet(doc, "GET /  —  page d'accueil présentant la plateforme.",
               runs_bold=("GET /",))
    add_bullet(doc, "GET /formations  —  liste dynamique des formations.",
               runs_bold=("GET /formations",))
    add_bullet(doc, "GET /formations/create  —  formulaire d'ajout.",
               runs_bold=("GET /formations/create",))
    add_bullet(doc, "POST /formations  —  enregistrement avec validation serveur.",
               runs_bold=("POST /formations",))
    add_bullet(doc, "DELETE /formations/{id}  —  suppression d'une formation "
                    "(bouton sur chaque carte avec confirmation JS).",
               runs_bold=("DELETE /formations/{id}",))

    add_heading(doc, "2.2 Diagramme des composants", level=2)
    add_code_block(doc, """\
                ┌──────────────────────────────────────────┐
                │              Docker host                 │
                │                                          │
  ┌─────────┐   │   ┌─────────────────┐   TCP  ┌────────┐  │
  │ Browser │──HTTP▶│  Service "app"  │───────▶│   db   │  │
  │         │ :8000 │  Laravel 10     │  :3306 │ MySQL8 │  │
  └─────────┘   │   │  PHP 8.2 + serve│        │        │  │
                │   └─────────────────┘        └────────┘  │
                │        network default (bridge)          │
                │                                          │
                │   volume : mysql_data (persistance DB)   │
                └──────────────────────────────────────────┘""")

    add_heading(doc, "2.3 Composants techniques", level=2)

    tbl = doc.add_table(rows=4, cols=3)
    tbl.style = "Light Grid Accent 1"
    tbl.autofit = True
    hdr = tbl.rows[0].cells
    for cell, text in zip(hdr, ["Composant", "Rôle", "Image / Base"]):
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    rows = [
        ("app",
         "Sert l'application Laravel via php artisan serve (port 8000)",
         "php:8.2-cli + Composer + extensions (pdo_mysql, mbstring, zip, bcmath)"),
        ("db",
         "Stocke les formations dans la table formations",
         "mysql:8.0 avec healthcheck mysqladmin ping"),
        ("mysql_data",
         "Volume Docker pour la persistance des données",
         "Volume nommé (driver local)"),
    ]
    for i, (a, b, c) in enumerate(rows, start=1):
        for j, text in enumerate((a, b, c)):
            tbl.rows[i].cells[j].text = text

    doc.add_paragraph()

    # =================================================================
    # 3. Captures d'écran
    # =================================================================
    doc.add_page_break()
    add_heading(doc, "3. Captures d'écran de l'application", level=1)

    add_heading(doc, "3.1 Page d'accueil (GET /)", level=2)
    add_paragraph(
        doc,
        "La page d'accueil présente la plateforme et propose un bouton "
        "d'appel à l'action vers la liste des formations.",
    )
    add_screenshot(doc, "01-home.png",
                   "Figure 1 — Page d'accueil de FormaPro")

    add_heading(doc, "3.2 Liste des formations (GET /formations)", level=2)
    add_paragraph(
        doc,
        "La liste affiche dynamiquement les formations enregistrées en base "
        "(6 formations seedées au premier démarrage), triées par niveau puis "
        "par titre. Chaque carte montre le titre, la description, la durée, "
        "un badge coloré selon le niveau et un bouton Supprimer qui déclenche "
        "une confirmation JavaScript avant d'envoyer une requête "
        "DELETE /formations/{id}.",
    )
    add_screenshot(doc, "02-formations.png",
                   "Figure 2 — Liste des formations avec boutons d'ajout et suppression")

    add_heading(doc, "3.3 Formulaire d'ajout (GET /formations/create)", level=2)
    add_paragraph(
        doc,
        "Le formulaire permet d'ajouter une nouvelle formation. La validation "
        "serveur vérifie que tous les champs sont remplis, que le titre est "
        "unique, que la description fait au moins 10 caractères et que le "
        "niveau appartient à la liste fermée (Débutant, Intermédiaire, Avancé).",
    )
    add_screenshot(doc, "03-create-form.png",
                   "Figure 3 — Formulaire d'ajout d'une formation")

    # =================================================================
    # 4. Automatisation
    # =================================================================
    doc.add_page_break()
    add_heading(doc, "4. Étapes d'automatisation", level=1)

    add_heading(doc, "4.1 Bootstrap sans installation globale", level=2)
    add_paragraph(
        doc,
        "Le squelette Laravel a été généré via un conteneur Composer jetable, "
        "sans avoir à installer PHP ou Composer sur la machine hôte :",
    )
    add_code_block(doc, 'docker run --rm -v "$PWD":/app -w /app composer:2 \\\n'
                         '    create-project laravel/laravel:^10.0 .')

    add_heading(doc, "4.2 Conteneurisation (Dockerfile)", level=2)
    add_paragraph(
        doc,
        "Le Dockerfile construit une image mono-processus à partir de "
        "php:8.2-cli, installe les extensions PHP requises, copie Composer "
        "depuis l'image officielle, puis installe les dépendances avec un "
        "cache Docker optimisé (composer.json/lock copiés avant le reste) :",
    )
    add_code_block(doc, """\
FROM php:8.2-cli

RUN apt-get update && apt-get install -y --no-install-recommends \\
        libzip-dev libpng-dev libonig-dev default-mysql-client \\
        zip unzip git \\
    && docker-php-ext-install pdo_mysql mbstring zip bcmath \\
    && rm -rf /var/lib/apt/lists/*

COPY --from=composer:2 /usr/bin/composer /usr/bin/composer

WORKDIR /var/www/html

COPY composer.json composer.lock ./
RUN composer install --no-dev --no-scripts --no-autoloader \\
    --prefer-dist --no-interaction

COPY . .
RUN composer dump-autoload --optimize --no-dev \\
    && chmod -R 775 storage bootstrap/cache

COPY docker/entrypoint.sh /usr/local/bin/entrypoint.sh
RUN chmod +x /usr/local/bin/entrypoint.sh

EXPOSE 8000
ENTRYPOINT ["entrypoint.sh"]
CMD ["php", "artisan", "serve", "--host=0.0.0.0", "--port=8000"]""")

    add_heading(doc, "4.3 Entrypoint intelligent", level=2)
    add_paragraph(
        doc,
        "Le script docker/entrypoint.sh rend le premier démarrage totalement "
        "automatique :",
    )
    add_bullet(doc, "Attend que MySQL soit prêt (via PDO — plus fiable que "
                    "mysqladmin avec MySQL 8 et TLS)")
    add_bullet(doc, "Copie .env.example vers .env si absent")
    add_bullet(doc, "Génère APP_KEY automatiquement")
    add_bullet(doc, "Exécute les migrations et seeders (--force --seed)")
    add_bullet(doc, "Cache la configuration et les routes")
    add_bullet(doc, 'Démarre php artisan serve via exec "$@"')

    add_heading(doc, "4.4 Orchestration (docker-compose.yml)", level=2)
    add_paragraph(
        doc,
        "Le fichier docker-compose.yml déclare les services app et db, "
        "avec un healthcheck MySQL consommé par "
        "depends_on: condition: service_healthy pour garantir l'ordre "
        "de démarrage.",
    )
    add_code_block(doc, """\
services:
  app:
    build: .
    image: formation-app:latest
    ports: ["8000:8000"]
    environment:
      DB_HOST: db
      DB_PORT: 3306
      DB_DATABASE: formations
      DB_USERNAME: laravel
      DB_PASSWORD: secret
      # ... (APP_NAME, APP_ENV, etc.)
    depends_on:
      db:
        condition: service_healthy

  db:
    image: mysql:8.0
    environment:
      MYSQL_DATABASE: formations
      MYSQL_USER: laravel
      MYSQL_PASSWORD: secret
      MYSQL_ROOT_PASSWORD: root
    volumes: [mysql_data:/var/lib/mysql]
    healthcheck:
      test: ["CMD", "mysqladmin", "ping", "-h", "localhost",
             "-u", "root", "-proot"]
      interval: 5s
      retries: 20

volumes:
  mysql_data:""")

    add_paragraph(doc, "État des conteneurs après docker compose up -d :", bold=True)
    add_code_block(doc, """\
NAME            IMAGE                  STATUS                  PORTS
formation-app   formation-app:latest   Up 3 hours              0.0.0.0:8000->8000/tcp
formation-db    mysql:8.0              Up 13 hours (healthy)   0.0.0.0:3306->3306/tcp""")

    add_heading(doc, "4.5 Pipeline CI/CD (GitHub Actions)", level=2)
    add_paragraph(
        doc,
        "Le workflow .github/workflows/ci.yml enchaîne trois jobs séquentiels :",
    )
    add_bullet(doc, "test  —  Installe Composer, exécute php artisan test sur "
                    "une base SQLite éphémère (pas de conteneur MySQL dans le "
                    "runner, ce qui garde le job rapide).",
               runs_bold=("test",))
    add_bullet(doc, "docker-build  —  Build de l'image via docker/build-push-action "
                    "avec cache GHA, puis sanity check (php artisan --version).",
               runs_bold=("docker-build",))
    add_bullet(doc, "deploy-stack  —  docker compose up -d --build dans le runner, "
                    "attente de disponibilité HTTP, smoke test sur / et /formations.",
               runs_bold=("deploy-stack",))
    add_paragraph(
        doc,
        "Le pipeline se déclenche sur push et pull_request des branches main/master, "
        "ainsi que manuellement via workflow_dispatch.",
    )

    # =================================================================
    # 5. Difficultés rencontrées
    # =================================================================
    doc.add_page_break()
    add_heading(doc, "5. Difficultés rencontrées", level=1)

    difficultes = [
        ("5.1 Chemins Windows et Git Bash",
         "Le bootstrap docker run -v \"$PWD\":/app a échoué sous Git Bash car MSYS "
         "convertit /app en C:/Program Files/Git/app. Solution : MSYS_NO_PATHCONV=1 "
         "et $(pwd -W) pour obtenir le chemin Windows natif, puis //app pour "
         "empêcher la conversion du chemin cible."),
        ("5.2 composer.lock incompatible PHP 8.2",
         "L'image officielle composer:2 embarque PHP 8.4, ce qui a généré un lock "
         "avec symfony/css-selector v8.0.8 (exigeant PHP ≥ 8.4). Le build dans "
         "php:8.2-cli plantait. Solution : ajouter config.platform.php = \"8.2\" "
         "dans composer.json puis relancer composer update pour régénérer un lock "
         "compatible (css-selector rétrogradé en v7.4.8)."),
        ("5.3 Ordre de démarrage MySQL / Laravel",
         "Au premier run, Laravel tentait de migrer avant que MySQL ait fini "
         "son initialisation. Deux niveaux de protection : healthcheck dans "
         "docker-compose.yml + boucle d'attente PHP (PDO::__construct) dans "
         "entrypoint.sh. À noter : le mysqladmin de Debian (fourni par MariaDB) "
         "ne gère pas bien le TLS de MySQL 8 — le client PDO PHP fonctionne "
         "sans configuration TLS particulière."),
        ("5.4 APP_KEY écrasée par l'environnement",
         "La première version du compose utilisait env_file: .env, qui injectait "
         "APP_KEY= (vide) dans l'environnement. Laravel lit env() depuis $_ENV "
         "avant le fichier .env, donc même après key:generate la clé était "
         "ignorée → 500 \"No application encryption key has been specified\". "
         "Solution : supprimer env_file:, déclarer uniquement les variables DB "
         "nécessaires dans environment:, laisser l'entrypoint générer APP_KEY "
         "dans le .env interne au conteneur."),
        ("5.5 Validation de niveau avec caractères accentués",
         "La règle de validation Laravel in:Débutant,Intermédiaire,Avancé "
         "a posé des soucis de parsing avec les caractères multibyte. "
         "Remplacement par Rule::in([...]) qui accepte directement un array "
         "— plus robuste, indépendant du parsing de la chaîne."),
    ]
    for titre, corps in difficultes:
        add_heading(doc, titre, level=2)
        add_paragraph(doc, corps)

    # =================================================================
    # 6. Bilan
    # =================================================================
    doc.add_page_break()
    add_heading(doc, "6. Bilan", level=1)

    add_heading(doc, "6.1 Récapitulatif des livrables", level=2)
    tbl = doc.add_table(rows=6, cols=2)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for cell, text in zip(hdr, ["Livrable", "Statut"]):
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    livrables = [
        ("Application fonctionnelle (accueil + liste + ajout + suppression)", "✅ Validé"),
        ("Dockerfile + build", "✅ Validé"),
        ("docker-compose.yml (app + mysql + healthcheck)", "✅ Validé"),
        ("Pipeline CI/CD (3 jobs : test, build, integration)", "✅ Validé"),
        ("Rapport avec architecture, étapes et difficultés", "✅ Ce document"),
    ]
    for i, (nom, statut) in enumerate(livrables, start=1):
        tbl.rows[i].cells[0].text = nom
        tbl.rows[i].cells[1].text = statut

    doc.add_paragraph()

    add_heading(doc, "6.2 Points forts", level=2)
    add_bullet(doc, "Zéro installation globale : tout passe par des conteneurs Docker "
                    "jetables (Composer, Python, Chromium).")
    add_bullet(doc, 'Premier démarrage "one-shot" : cp .env.example .env && '
                    'docker compose up -d --build suffit.')
    add_bullet(doc, "Pipeline à trois étages qui valide code → image → stack complète.")
    add_bullet(doc, "Application étendue avec un formulaire d'ajout (GET + POST "
                    "avec validation serveur) et une suppression (DELETE avec "
                    "confirmation JS), chacune suivie d'un message flash.")

    add_heading(doc, "6.3 Améliorations possibles", level=2)
    add_bullet(doc, "Remplacer php artisan serve par Nginx + PHP-FPM pour un "
                    "déploiement plus proche de la production.")
    add_bullet(doc, "Ajouter un reverse-proxy Traefik avec HTTPS automatique.")
    add_bullet(doc, "Publier l'image sur Docker Hub ou GHCR depuis le pipeline.")
    add_bullet(doc, "Étoffer les tests Feature (actuellement seulement le test "
                    "par défaut de Laravel).")

    add_heading(doc, "6.4 Lien vers le code", level=2)
    add_paragraph(
        doc,
        "Le code source complet est disponible sur GitHub : "
        "https://github.com/GhostSt94/mini-projet-automatisation",
    )

    # =================================================================
    # Save
    # =================================================================
    doc.save(OUT)
    print(f"[OK] Rapport généré : {OUT}")
    print(f"     Taille : {OUT.stat().st_size / 1024:.1f} Kio")


if __name__ == "__main__":
    build()
